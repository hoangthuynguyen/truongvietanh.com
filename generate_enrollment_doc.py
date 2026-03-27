#!/usr/bin/env python3
"""
Generate comprehensive DOCX for Trường Việt Anh 2026 Enrollment Operating System
Version 11 - Complete Vietnamese document with 19 chapters
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add header
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.text = "Kế Hoạch Tuyển Sinh 2026 - V11 Enrollment Operating System"
header_para.style = 'Normal'
for run in header_para.runs:
    run.font.size = Pt(10)
    run.font.italic = True

# Add footer with page numbers
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar1)
run = footer_para.add_run()
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "PAGE"
run._r.append(instrText)
run = footer_para.add_run()
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar2)

def add_heading(text, level=1):
    """Add a heading with appropriate style"""
    if level == 1:
        p = doc.add_heading(text, level=1)
    elif level == 2:
        p = doc.add_heading(text, level=2)
    else:
        p = doc.add_heading(text, level=3)
    return p

def add_paragraph(text, bold=False, space_before=6, space_after=6):
    """Add a paragraph with formatting"""
    p = doc.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    if bold:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(12)
    else:
        for run in p.runs:
            run.font.size = Pt(11)
    return p

def add_bullet_list(items):
    """Add bulleted list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)

def add_table(rows, col_widths=None):
    """Add table with data"""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)

    # Populate table
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(cell_data)

            # Format header row
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '4472C4')
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    return table

def add_page_break():
    """Add page break"""
    doc.add_page_break()

# ============= TITLE PAGE =============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('KỂ HOẠCH TUYỂN SINH 2026')
title_run.font.size = Pt(28)
title_run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Phiên Bản 11 - Enrollment Operating System')
subtitle_run.font.size = Pt(16)
subtitle_run.bold = True

school = doc.add_paragraph()
school.alignment = WD_ALIGN_PARAGRAPH.CENTER
school.paragraph_format.space_before = Pt(24)
school_run = school.add_run('Trường Việt Anh')
school_run.font.size = Pt(14)
school_run.bold = True

year = doc.add_paragraph()
year.alignment = WD_ALIGN_PARAGRAPH.CENTER
year.paragraph_format.space_before = Pt(24)
year_run = year.add_run('2026')
year_run.font.size = Pt(12)

add_page_break()

# ============= TABLE OF CONTENTS =============
add_heading('MỤC LỤC', 1)
toc_items = [
    'PHẦN A - NỀN TẢNG HỆ THỐNG',
    '  Chương 1: Tổng Quan Enrollment Operating System',
    '  Chương 2: Chân Dung Khách Hàng Mục Tiêu',
    '  Chương 3: Funnel 8 Tầng - Pipeline Tuyển Sinh',
    '  Chương 4: SLA Tuyển Sinh - Cam Kết Dịch Vụ',
    '  Chương 5: SOP School Tour - Quy Trình Tham Quan Trường',
    '',
    'PHẦN B - CẤU TRÚC KÊNH & NGÂN SÁCH',
    '  Chương 6: Cấu Trúc Kênh 4 Trục',
    '  Chương 7: Budget Governance - Quản Trị Ngân Sách',
    '  Chương 8: Chi Tiết Paid Media Theo Kênh',
    '',
    'PHẦN C - VẬN HÀNH HÀNG NGÀY',
    '  Chương 9: Dashboard 12 Số Hàng Ngày',
    '  Chương 10: Weekly Battleplan - 5 Phases',
    '  Chương 11: Campus Ownership & RACI',
    '  Chương 12: Content Engine - Nội Dung Theo Funnel',
    '',
    'PHẦN D - CÔNG NGHỆ & TRACKING',
    '  Chương 13: Hệ Thống CRM Kép (GHL + Pancake)',
    '  Chương 14: Funnel Chi Tiết & Landing Pages',
    '  Chương 15: Analytics, GTM & Remarketing',
    '',
    'PHẦN E - KIỂM SOÁT & CẢI TIẾN',
    '  Chương 16: QA & Mystery Call Program',
    '  Chương 17: Lost Reasons Tracking & Win-Back',
    '  Chương 18: Offer & Deadline Management',
    '  Chương 19: Nhân Sự & Đào Tạo'
]
for item in toc_items:
    p = doc.add_paragraph(item)
    if item.startswith('  '):
        p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.size = Pt(10)

add_page_break()

# ============= PHẦN A - NỀN TẢNG HỆ THỐNG =============
add_heading('PHẦN A - NỀN TẢNG HỆ THỐNG', 1)
add_page_break()

# ============= CHAPTER 1 =============
add_heading('Chương 1: Tổng Quan Enrollment Operating System', 1)
add_paragraph('Tầm nhìn chuyên nghiệp: Chuyển đổi từ "chạy quảng cáo tuyển sinh" sang "vận hành hệ thống tuyển sinh chuyên nghiệp" với quy trình, số liệu, và kiểm soát.')
add_page_break()

add_heading('1.1 Các KPI Tổng Thể', 2)
add_table([
    ['Chỉ Tiêu', 'Giá Trị', 'Ghi Chú'],
    ['Mục tiêu tuyển sinh', '1,300 học sinh', 'Toàn hệ thống 6 campus'],
    ['Ngân sách', '13 tỷ VND', '10.8 tỷ cam kết + 2.2 tỷ linh hoạt'],
    ['Thời gian', '28 tuần (26/03 - 30/09/2026)', 'Từ Warm-up đến Close'],
    ['CAC mục tiêu', '< 10 triệu/học sinh', 'Điều chỉnh theo cấp học'],
    ['Conversion rate', '4.5% từ lead → enroll', 'Dựa trên funnel benchmarks']
], [1.5, 2.0, 2.5])

add_heading('1.2 Timeline 4 Giai Đoạn', 2)
add_paragraph('Chiến dịch chia thành 4 giai đoạn chiến lược:')
add_bullet_list([
    'Warm-up (T3-T4 | 26/03-30/04): Xây dựng brand awareness, launch landing pages, kích hoạt retargeting audiences',
    'Ramp-up (T5-T6 | 01/05-30/06): Tăng budget chi trả, optimize campaigns, launch summer programs',
    'Peak (T7-T8 | 01/07-31/08): Cuộc tổng tấn công đa kênh, peak demand, aggressive closing',
    'Close (T9 | 01/09-30/09): Nurture đơn chưa hoàn thành, last-minute offers, deadline urgency'
])

add_heading('1.3 20 Quyết Định Phải Lock Trong 72 Giờ Đầu Tiên', 2)
add_bullet_list([
    'Approval ngân sách committed 10.8 tỷ',
    'Phân bổ budget chi tiết theo 4 giai đoạn và 6 kênh',
    'Lộ trình landing page hoàn thành (deadline mỗi trang)',
    'Danh sách keywords Google Ads finalize',
    'Creative assets (video, image, copy) upload & approve',
    'Influencer partners & collaboration timeline confirm',
    'Campus leads assign cho mỗi cơ sở',
    'Admissions team training schedule',
    'Open Day dates & agendas finalize',
    'Webinar topics & speakers lock',
    'Telesales scripts & objection handling approve',
    'CRM workflows setup & test (GHL + Pancake)',
    'Make.com sync workflows test & live',
    'Analytics setup (GTM tags, CAPI, pixel)',
    'Remarketing audiences create & test',
    'Email sequences write & approve',
    'SMS templates finalize',
    'Zalo OA content calendar upload',
    'Mystery call process define & schedule',
    'Weekly meeting agenda & participants confirm'
])

add_page_break()

# ============= CHAPTER 2 =============
add_heading('Chương 2: Chân Dung Khách Hàng Mục Tiêu (4 Personas)', 1)
add_paragraph('Cơ sở dữ liệu khách hàng mục tiêu gồm 4 personas theo cấp học, mỗi persona có pain points, trigger events, decision factors riêng biệt.')

add_heading('2.1 Persona 1: Ba Mẹ Mầm Non (25-35 tuổi)', 2)
add_table([
    ['Khía Cạnh', 'Chi Tiết'],
    ['Độ tuổi con', '2-5 tuổi'],
    ['Trình độ giáo dục', 'Đại học trở lên (70%)'],
    ['Thu nhập', '15-40 triệu/tháng'],
    ['Lối sống', 'Đô thị, chủ động tìm tòi'],
    ['Pain points', 'Chất lượng giáo dục sớm, tiếng Anh, bạn bè đa dạng'],
    ['Trigger event', 'Con chuẩn bị vào MN, chuyển nhà mới'],
    ['Decision factors', 'Chương trình giáo dục, cơ sở vật chất, biểu phí'],
    ['Information sources', 'Facebook, Google, recommendation, school website'],
    ['Decision timeline', '30-45 ngày từ awareness → enroll'],
    ['Budget concerns', 'Vẫn cân nhắc, cần justify ROI'],
    ['Geographic focus', 'Gò Vấp, Bình Tân, Tây Ninh']
], [2.0, 4.0])

add_heading('2.2 Persona 2: Ba Mẹ Tiểu Học (30-40 tuổi)', 2)
add_table([
    ['Khía Cạnh', 'Chi Tiết'],
    ['Độ tuổi con', '6-11 tuổi'],
    ['Trình độ giáo dục', 'Đại học - Cao học (75%)'],
    ['Thu nhập', '20-50 triệu/tháng'],
    ['Lối sống', 'Bận rộn, cần school quality & convenience'],
    ['Pain points', 'Tiếng Anh + Toán, extra-curricular activities, discipline'],
    ['Trigger event', 'Hết cấp MN, không hài lòng school cũ'],
    ['Decision factors', 'Academic reputation, teacher quality, program breadth'],
    ['Information sources', 'Facebook group, Google review, parent network'],
    ['Decision timeline', '45-60 ngày từ awareness → enroll'],
    ['Budget concerns', 'Đã quyết định chi tiêu, cần good value'],
    ['Geographic focus', 'Gò Vấp (priority), Bình Tân']
], [2.0, 4.0])

add_heading('2.3 Persona 3: Ba Mẹ Cấp 2 (35-45 tuổi)', 2)
add_table([
    ['Khía Cạnh', 'Chi Tiết'],
    ['Độ tuổi con', '12-14 tuổi'],
    ['Trình độ giáo dục', 'Cao học (60%), Kinh doanh chủ động'],
    ['Thu nhập', '30-60+ triệu/tháng'],
    ['Lối sống', 'International mindset, cẩn thận trong chọn lựa'],
    ['Pain points', 'Academic excellence, preparation cho THPT/International, tạo hạnh nhân'],
    ['Trigger event', 'Hết cấp TH, chuẩn bị exam entrance, so sánh schools'],
    ['Decision factors', 'Curriculum quality, university prep, facilities, reputation'],
    ['Information sources', 'Google search, school website, parent review, tour'],
    ['Decision timeline', '60-75 ngày từ awareness → enroll'],
    ['Budget concerns', 'Ít lo lắng, focus quality'],
    ['Geographic focus', 'Gò Vấp, Bình Tân, vilanh từ tỉnh']
], [2.0, 4.0])

add_heading('2.4 Persona 4: Ba Mẹ Cấp 3 (38-50 tuổi)', 2)
add_table([
    ['Khía Cạnh', 'Chi Tiết'],
    ['Độ tuổi con', '15-18 tuổi'],
    ['Trình độ giáo dục', 'Cao học (70%), Leadership roles'],
    ['Thu nhập', '40-80+ triệu/tháng'],
    ['Lối sống', 'Global vision, strategic planning, luxury-conscious'],
    ['Pain points', 'University admission success, international pathway, character development'],
    ['Trigger event', 'Hết cấp THCS, chuẩn bị Đại học, cần advanced curriculum'],
    ['Decision factors', 'International pathways, university outcomes, pedagogy innovation'],
    ['Information sources', 'School website, direct inquiry, top-tier groups, expert advice'],
    ['Decision timeline', '75-90 ngày từ awareness → enroll (can be 180+ for hesitant)'],
    ['Budget concerns', 'Premium positioning, value story critical'],
    ['Geographic focus', 'Gò Vấp (core), suburban HCM, business districts']
], [2.0, 4.0])

add_page_break()

# ============= CHAPTER 3 =============
add_heading('Chương 3: Funnel 8 Tầng - Pipeline Tuyển Sinh', 1)
add_paragraph('Mô hình pipeline tuyển sinh có 8 tầng rõ ràng, mỗi tầng có tiêu chí chuyển đổi cụ thể.')

add_heading('3.1 Định Nghĩa 8 Tầng', 2)
add_bullet_list([
    'New Lead: Dữ liệu mới từ ads/organic/referral, chưa contacted',
    'Contacted: Đã gọi/message, có phản hồi hoặc message pending reply',
    'Qualified: Xác nhận nhu cầu, budget, timeline; sẵn sàng booking',
    'Booked Tour: Đã đặt hẹn tham quan, confirmation sent',
    'Showed Tour: Đã tham quan trực tiếp trường',
    'Offer: Tư vấn viên đã đưa lời mời (quote, pricing, program)',
    'Deposit: Gia đình đã thanh toán deposit/tiền nạp (confirm intent)',
    'Enrolled: Hoàn tất enrollment, sẵn sàng nhập học'
])

add_heading('3.2 Benchmark Conversion Rates', 2)
add_table([
    ['Funnel Stage', 'Target Conversion %', 'Meaning'],
    ['Lead → Contacted', '90%', '9 trong 10 lead được contact'],
    ['Contacted → Qualified', '60%', '6 trong 10 contacted thành qualified'],
    ['Qualified → Booked Tour', '50%', '5 trong 10 qualified đặt hẹn'],
    ['Booked → Showed Tour', '70%', '7 trong 10 booked thực sự đến'],
    ['Showed → Offer', '80%', '8 trong 10 đã tour nhận offer'],
    ['Offer → Deposit', '60%', '6 trong 10 offer convert → deposit'],
    ['Deposit → Enrolled', '95%', '95% deposit đủ điều kiện nhập học']
], [2.5, 2.5, 1.5])

add_heading('3.3 Pipeline Math - Bao Nhiêu Leads Cần Ở Mỗi Tầng?', 2)
add_paragraph('Công thức tính ngược từ mục tiêu 1,300 enrolled:')
add_bullet_list([
    '1,300 enrolled ÷ 95% (Deposit→Enrolled) = 1,368 deposits cần',
    '1,368 deposits ÷ 60% (Offer→Deposit) = 2,280 offers cần',
    '2,280 offers ÷ 80% (Showed→Offer) = 2,850 showed tour cần',
    '2,850 showed ÷ 70% (Booked→Showed) = 4,071 booked hẹn cần',
    '4,071 booked ÷ 50% (Qualified→Booked) = 8,143 qualified cần',
    '8,143 qualified ÷ 60% (Contacted→Qualified) = 13,572 contacted cần',
    '13,572 contacted ÷ 90% (Lead→Contacted) = 15,080 leads cần'
])
add_paragraph('Kết luận: Cần ~15,080 leads (mới) để đạt 1,300 enrolled với benchmarks này.', bold=True)

add_page_break()

# ============= CHAPTER 4 =============
add_heading('Chương 4: SLA Tuyển Sinh - Cam Kết Dịch Vụ', 1)
add_paragraph('SLA định rõ response times, touchpoints, recovery procedures để bảo đảm không lead nào bị "rơi thuyền".')

add_heading('4.1 Response Time Targets', 2)
add_table([
    ['Loại Hình', 'Trong Giờ Làm Việc', 'Ngoài Giờ'],
    ['Facebook Lead', '3 phút', '15 phút'],
    ['Google Lead', '5 phút', '20 phút'],
    ['Walk-in/Visit', '2 phút (tiếp đón)', 'N/A'],
    ['Phone Call', 'Pick up < 1 phút', 'Callback < 15 phút'],
    ['Email', '2 giờ', '4 giờ (next business day)'],
    ['SMS Inquiry', '5 phút', '15 phút']
], [2.0, 2.0, 2.0])

add_heading('4.2 Touchpoint Requirement (7 Ngày Đầu)', 2)
add_paragraph('Mỗi lead mới phải nhận minimum 8-10 touchpoints trong 7 ngày đầu:')
add_bullet_list([
    'Touchpoint 1: Initial response (call, message, email)',
    'Touchpoint 2-3: Nurture messages (2-3 touchpoints via Zalo/SMS/email)',
    'Touchpoint 4: Schedule tour (explicit booking confirmation)',
    'Touchpoint 5: Pre-tour reminder (12 giờ trước)',
    'Touchpoint 6: Tour itself',
    'Touchpoint 7-8: Post-tour follow-up (1 hari & 3 ngày)',
    'Touchpoint 9-10: Offer + Deadline reminder'
])

add_heading('4.3 No-Show Recovery', 2)
add_paragraph('Nếu lead không đến tour:')
add_bullet_list([
    'Step 1: Gọi lại trong 30 phút sau no-show (xác nhận lý do)',
    'Step 2: Offer reschedule với flexibility (khác ngày, khác giờ)',
    'Step 3: Virtual tour option nếu lead ở xa',
    'Step 4: After 2 reschedule fails → move to "Lost → Recovery" campaign'
])

add_heading('4.4 Escalation Process', 2)
add_bullet_list([
    'Miss SLA 1x → Quản lý team lên tiếng cải thiện',
    'Miss SLA 2-3 lần trong tuần → Meeting với Campus Lead',
    'Miss SLA liên tục → Performance review, retraining, potential replacement'
])

add_page_break()

# ============= CHAPTER 5 =============
add_heading('Chương 5: SOP School Tour - Quy Trình Tham Quan Trường (75 Phút)', 1)
add_paragraph('Quy trình tham quan được tiêu chuẩn hóa để bảo đảm experience nhất quán & chuyên nghiệp, tăng tour→offer conversion.')

add_heading('5.1 Timeline Cấu Trúc (75 Phút)', 2)
add_table([
    ['Segment', 'Thời Gian', 'Chi Tiết'],
    ['1. Đón Tiếp & Ấm Cúng', '10 phút', 'Đặt vị trí, cà phê, hỏi nhu cầu'],
    ['2. Presentation', '15 phút', 'Giá trị cốt lõi, khác biệt, curriculum'],
    ['3. Tour Cơ Sở Vật Chất', '25 phút', 'Route map, highlight key facilities'],
    ['4. Q&A & Tư Vấn', '15 phút', 'Match nhu cầu → solution, resolve concerns'],
    ['5. Offer & Next Steps', '10 phút', 'Present pricing, tạo urgency, đặt deadline']
], [1.5, 1.5, 3.0])

add_heading('5.2 Segment 1: Đón Tiếp (10 Phút)', 2)
add_paragraph('Script: "Chào mừng gia đình đến Trường Việt Anh. Xin vui lòng ngồi, tôi sẽ mang cà phê. Hôm nay gia đình quan tâm đặc biệt cấp nào & program nào?"')
add_bullet_list([
    'Chào hỏi thân thiện, giới thiệu self',
    'Offer nước/cà phê (tạo comfort)',
    'Hỏi các câu open-end: Tại sao gia đình chọn VA? Nhu cầu chính là gì?',
    'Ghi chú chi tiết vào CRM',
    'Set expectation: "Tour sẽ mất 75 phút, sau tour mình có lời đề cử cho gia đình"'
])

add_heading('5.3 Segment 2: Presentation (15 Phút)', 2)
add_paragraph('Trình bày core value proposition:')
add_bullet_list([
    'Trường Việt Anh là: International standard education, local values, 6 campuses established',
    'Curriculum highlight: Bilingual, STEM-focused, character development',
    'Outcomes: 99% university admission, international pathway',
    'Use slides + printed materials (không quá dài)',
    'Pause để Q&A tinh tươi'
])

add_heading('5.4 Segment 3: Tour Cơ Sở Vật Chất (25 Phút)', 2)
add_paragraph('Route map per campus (ví dụ Gò Vấp):')
add_bullet_list([
    'Classrooms (khám phá learning environment)',
    'Science lab (hands-on STEM)',
    'Library (resources, quiet space)',
    'Playground/Sports (facilities)',
    'Cafeteria (nutrition, safety)',
    "Teacher's lounge (expert staff)",
    'Admin office (tổ chức hành chính)'
])
add_paragraph('Highlight points: Safety features, technology integration, student feedback posted.')

add_heading('5.5 Segment 4: Q&A & Tư Vấn (15 Phút)', 2)
add_paragraph('Xử lý objections phổ biến:')
add_table([
    ['Objection', 'Response Script'],
    ['"Giá cao quá"', '"Chất lượng giáo dục Quốc tế tiêu chuẩn cao, campus hạng A, teacher quality là top. Giáo dục con em là đầu tư dài hạn nhất."'],
    ['"Xa nhà"', '"Có các campus ở Bình Tân, Tây Ninh gần hơn. Ngoài ra shuttle service, carpool community quanh campus."'],
    ['"Chưa quyết định"', '"Hoàn toàn bình thường. Gia đình cân nhắc, so sánh khác school. Đó là lý do tôi tặng early bird pricing & limited spots offer hôm nay."'],
    ['"So sánh khác trường"', '"Rất tự tin so sánh. Mời gia đình tour schools khác, sau đó compare. Kami confident gia đình chọn VA."']
], [2.0, 4.0])

add_heading('5.6 Segment 5: Offer & Next Steps (10 Phút)', 2)
add_paragraph('Create urgency & close:')
add_bullet_list([
    'Present offer: "Gia đình đăng ký hôm nay, nhận early bird 10% tuition discount, valid đến [DATE]"',
    'Show limited spots: "Chúng ta chỉ còn 8 spots cấp 1 lớp này, nếu gia đình muốn, cần confirm ngay"',
    'Next step: "Điền form đăng ký, deposit 1 triệu xác nhận, sau đó mình schedule placement test"',
    'Handover: "Tôi kết nối gia đình với Admissions Officer để hoàn tất paperwork"',
    'Call-to-action: "Có câu hỏi gì trước khi gia đình quyết định?"'
])

add_heading('5.7 Checklist Chuẩn Bị Vật Phẩm', 2)
add_bullet_list([
    'Folder brochure (đẹp, full color, info đầy đủ)',
    'Pricing sheet (rõ ràng, so sánh packages)',
    'Enrollment form (sẵn sàng cho signup)',
    'Small gift (pen, notepad, school merch)',
    'Name tag (for staff)',
    'Printed testimonials/parent quotes',
    'Campus photos/videos (if applicable)',
    'iPad/tablet (for digital presentation)',
    'CRM login (để update real-time)'
])

add_page_break()

# ============= PHẦN B =============
add_heading('PHẦN B - CẤU TRÚC KÊNH & NGÂN SÁCH', 1)
add_page_break()

# ============= CHAPTER 6 =============
add_heading('Chương 6: Cấu Trúc Kênh 4 Trục', 1)
add_paragraph('Chiến dịch tuyển sinh hoạt động trên 4 trục kênh, mỗi trục phục vụ mục tiêu khác nhau trong customer journey.')

add_heading('6.1 Trục 1 - Intent Capture (Google Ads, SEO)', 2)
add_paragraph('Bắt khách hàng đang tìm kiếm proactively (high intent):')
add_bullet_list([
    'Google Search Ads: keywords như "tuyển sinh cấp 1 HCM", "Trường quốc tế Gò Vấp", "School near Bình Tân"',
    'Google Display: banner ads on education websites, parent forums',
    'YouTube Pre-roll: educational content channels (parenting, education)',
    'SEO: Optimize website cho ranking keywords (organic cost-free)',
    'KPI: Lead cost < 1M VND/lead, CTR > 3%',
    'Budget: 30% of committed budget'
])

add_heading('6.2 Trục 2 - Demand Creation (Meta, TikTok, Zalo)', 2)
add_paragraph('Tạo nhu cầu mới, nurture potential customers:')
add_bullet_list([
    'Facebook Ads: Awareness campaigns (video, carousel), lookalike audiences',
    'Instagram Ads: Visual storytelling (campus life, student testimonials)',
    'TikTok Ads: Short-form video (viral potential, Gen Z parents targeting)',
    'Zalo Ads: Follower building, direct messaging',
    'KPI: CPM < 20K VND, Engagement rate > 2%',
    'Budget: 35% of committed budget'
])

add_heading('6.3 Trục 3 - Nurture (Email, SMS, Zalo OA, Retargeting)', 2)
add_paragraph('Nuôi dưỡng leads chưa sẵn sàng quyết định:')
add_bullet_list([
    'Email sequences: Welcome, value education, social proof, final deadline',
    'SMS: Reminders, urgency, booking confirmation',
    'Zalo OA: Rich content (video, testimonials, FAQs, 1-1 support)',
    'Retargeting: Dynamic ads targeting page visitors, abandoned inquiries',
    'KPI: Open rate > 25%, Click rate > 5%',
    'Budget: 20% of committed budget (mostly platform costs)'
])

add_heading('6.4 Trục 4 - Conversion (School Tour, Open Day, Webinar, Telesales)', 2)
add_paragraph('Đóng đơn, convert leads thành enrolled students:')
add_bullet_list([
    'School Tour: Scheduled campus visits (biggest conversion tool)',
    'Open Day Events: Large-scale events (multiple tours per day)',
    'Webinar: Expert panels (admission secrets, curriculum overview)',
    'Telesales: 1-1 phone outreach (handling objections, closing)',
    'KPI: Tour-to-offer > 80%, Offer-to-deposit > 60%',
    'Budget: 15% of committed budget'
])

add_page_break()

# ============= CHAPTER 7 =============
add_heading('Chương 7: Budget Governance - Quản Trị Ngân Sách', 1)
add_paragraph('Ngân sách 13 tỷ VND được quản lý theo khung Committed + Flexible, với CAC caps theo cấp học.')

add_heading('7.1 Committed Budget: 10.8 tỷ VND', 2)
add_table([
    ['Giai Đoạn', 'Tháng', 'Budget (tỷ VND)', 'Allocation %'],
    ['Warm-up', 'T3-T4 (8 tuần)', '2.0', '18.5%'],
    ['Ramp-up', 'T5-T6 (8 tuần)', '3.5', '32.4%'],
    ['Peak', 'T7-T8 (8 tuần)', '4.0', '37.0%'],
    ['Close', 'T9 (4 tuần)', '1.3', '12.0%'],
    ['TOTAL', '', '10.8', '100%']
], [1.5, 1.5, 1.5, 1.5])

add_heading('7.2 Flexible Budget: 2.2 tỷ VND (Release Conditions)', 2)
add_paragraph('Flexible budget được release khi kênh đạt target CPA:')
add_bullet_list([
    'Condition: Kênh nào đạt CPA < target sau 7 ngày → release 20% flexible budget cho kênh đó',
    'Escalation: Nếu CPA đang 10% dưới target → suggest double the spend',
    'Monitoring: Daily CPA tracking, weekly allocation rebalance',
    'Holdback: 30% flexible budget giữ lại cho last-minute campaigns (tháng 9)'
])

add_heading('7.3 CAC Caps Per Grade Level', 2)
add_table([
    ['Cấp Học', 'Mục Tiêu CAC', 'Rationale'],
    ['MN (Mầm non)', '< 8 triệu', 'Volume cao, conversion dễ (ít competitors)'],
    ['TH (Tiểu học)', '< 10 triệu', 'Volume trung bình, competitive'],
    ['THCS (Cấp 2)', '< 12 triệu', 'Volume thấp, high-decision parents'],
    ['THPT (Cấp 3)', '< 15 triệu', 'Rarest (existing students skip to), premium']
], [2.0, 2.0, 2.0])

add_heading('7.4 Kill Switch Rules', 2)
add_bullet_list([
    'Rule: Nếu campaign CPA > 2x target CAC sau 7 ngày liên tục → tắt campaign',
    'Exception: Chỉ turn off nếu đã tried 2-3 optimization attempts (creative swap, audience refine)',
    'Recovery: Learning phase tối đa 10 ngày, sau đó evaluate và cut losses',
    'Alert: Auto-alert nếu CPA trending upward > 15% 3 ngày liên tiếp'
])

add_heading('7.5 Monthly Budget Review Process', 2)
add_bullet_list([
    'Every Friday 2pm: Weekly finance check (spend vs. budget, CPA vs. target)',
    'Every Month 1st: Full budget review (month performance, next month reallocation)',
    'Attendees: Marketing Manager, Finance, Principal, Campus Leads',
    'Output: Budget adjustment memo (signed approval required for >10% variance)'
])

add_page_break()

# ============= CHAPTER 8 =============
add_heading('Chương 8: Chi Tiết Paid Media Theo Kênh', 1)

add_heading('8.1 Google Ads Strategy', 2)
add_table([
    ['Campaign Type', 'Keywords', 'Target', 'Budget %'],
    ['Search - MN', 'mầm non HCM, preschool Gò Vấp, nhà trẻ quốc tế', 'Parents 25-35', '15%'],
    ['Search - TH', 'tiểu học Gò Vấp, trường cấp 1 tuyệt vời', 'Parents 30-40', '20%'],
    ['Search - THCS', 'THCS Gò Vấp, trường cấp 2 quốc tế', 'Parents 35-45', '15%'],
    ['Search - THPT', 'THPT quốc tế, cấp 3 HCM', 'Parents 38-50', '10%'],
    ['Display Network', 'Education site placements, parent forums', 'Broad audience', '20%'],
    ['YouTube', 'Pre-roll on education channels', 'Parents researching', '15%'],
    ['Performance Max', 'Auto-optimization, all Google surfaces', 'High intent', '5%']
], [1.5, 1.8, 1.2, 1.5])

add_heading('8.2 Meta Ads (Facebook + Instagram)', 2)
add_paragraph('Meta is the primary demand creation engine (35% of budget):')
add_bullet_list([
    'Lead Ads: Direct form submission (no landing page friction)',
    'Conversion Ads: Direct-to-tour booking or phone call',
    'Video Ads: 15-30s campus highlights, parent testimonials',
    'Carousel Ads: Program comparison (MN vs. TH, curriculum highlight)',
    'Retargeting: Website visitors, video viewers, lead database',
    'Lookalike Audiences: Copy successful lead source profiles',
    'Budget allocation: 20% awareness, 15% conversion, rest retargeting'
])

add_heading('8.3 TikTok Ads Strategy', 2)
add_bullet_list([
    'Spark Ads: Native TikTok content (less disruptive)',
    'In-feed Ads: Video within FYP feed (high engagement)',
    'Hashtag Challenge: User-generated content campaign (optional, high effort)',
    'Target: Gen X/Millennial parents (surprising TikTok user base)',
    'Creative: 15-30s vertical video, student testimonials, day-in-life campus',
    'Budget: 15% of Demand Creation, scaling if ROI positive'
])

add_heading('8.4 Zalo Ads & OA Strategy', 2)
add_bullet_list([
    'Zalo Ads: Zalo Official Account follower acquisition',
    'ZaloOA Content: Rich messages (video, image, PDF), 1-1 customer support',
    'Conversion Ad: Direct booking link, phone call CTA',
    'Segmentation: Followers tagged by grade (MN/TH/THCS/THPT)',
    'Automation: Welcome sequence, nurture drip, deadline reminder',
    'Budget: 10% of demand budget, high-margin channel'
])

add_heading('8.5 Monthly Budget Allocation per Channel', 2)
add_table([
    ['Channel', 'T3-T4', 'T5-T6', 'T7-T8', 'T9'],
    ['Google Ads (Search)', '300M', '600M', '700M', '200M'],
    ['Google Display/YouTube', '150M', '350M', '400M', '100M'],
    ['Meta (FB+IG)', '600M', '1,000M', '1,200M', '400M'],
    ['TikTok', '150M', '300M', '400M', '100M'],
    ['Zalo Ads', '100M', '200M', '250M', '50M'],
    ['Retargeting (Dynamic)', '200M', '400M', '500M', '150M'],
    ['Content/Influencer', '150M', '250M', '350M', '50M'],
    ['Reserve/Testing', '250M', '400M', '200M', '150M'],
    ['TOTAL', '2,000M', '3,500M', '4,000M', '1,300M']
], [1.2, 0.9, 0.9, 0.9, 1.1])

add_page_break()

# ============= PHẦN C =============
add_heading('PHẦN C - VẬN HÀNH HÀNG NGÀY', 1)
add_page_break()

# ============= CHAPTER 9 =============
add_heading('Chương 9: Dashboard 12 Số Hàng Ngày', 1)
add_paragraph('12 metrics được track real-time hàng ngày để giám sát sức khỏe chiến dịch.')

add_heading('9.1 Dashboard Daily 12 Metrics', 2)
add_table([
    ['#', 'Metric', 'Target Ngày', 'Data Source', 'Alert If'],
    ['1', 'Spend Hôm Nay', 'Theo plan (40M)', 'GHL + Birch', '> 50% deviation'],
    ['2', 'Leads Hôm Nay', 'Theo funnel (40)', 'GHL + LP', '< 30 leads'],
    ['3', 'Qualified Leads', 'Theo plan (24)', 'GHL status', '< 20 qual'],
    ['4', 'CPL', '< 300K', 'Spend / Leads', '> 350K'],
    ['5', 'Avg Response Time', '< 5 phút', 'GHL timestamp', '> 7 phút'],
    ['6', 'Tours Booked', 'Theo plan (10)', 'GHL pipeline', '< 5 tours'],
    ['7', 'Show-up Rate %', '> 70%', 'Attended / Booked', '< 60%'],
    ['8', 'Offers Made', 'Theo plan (8)', 'GHL offer stage', '< 5 offers'],
    ['9', 'Deposits Received', 'Theo plan (4)', 'GHL + Bank', '< 2 deps'],
    ['10', 'Enrollments Confirm', 'Theo plan (2)', 'GHL + Admin', '< 1 enroll'],
    ['11', 'CAC Running', '< Target per grade', 'Cumulative', '> CAC cap'],
    ['12', 'Forecast vs Gap', 'Track weekly', 'Math model', '> 10% gap']
], [0.4, 1.0, 1.2, 1.2, 1.2])

add_heading('9.2 Alert Triggers & Actions', 2)
add_bullet_list([
    'RED Alert (Immediate): Any metric > 30% off target → call Marketing Manager right away',
    'YELLOW Alert (Quick Review): Any metric > 15-30% off target → review in next 2 hours, assess',
    'GREEN (Normal): Within 15% → continue, monitor trend',
    'Trend Alert: Same metric yellow 3 days straight → escalate to Principal'
])

add_heading('9.3 Dashboard Tool & Setup', 2)
add_paragraph('Dashboard source: Google Sheet synced from GHL API + Birch API')
add_bullet_list([
    'Owner: Digital Specialist',
    'Update frequency: Auto-refresh hourly',
    'Access: Marketing Manager + Principal (read), Digital Specialist (edit)',
    'Backup: Daily CSV export, stored in Google Drive'
])

add_page_break()

# ============= CHAPTER 10 =============
add_heading('Chương 10: Weekly Battleplan - 5 Phases', 1)
add_paragraph('Tuần được chia 5 phase rõ ràng, mỗi phase có objective, attendees, deliverables.')

add_heading('10.1 Phase 1 - Monday Data Review (30 Phút, 9:00am)', 2)
add_bullet_list([
    'Attendees: Marketing Manager, Digital Specialist, Campus Leads',
    'Objective: Review 12 dashboard metrics from past week',
    'Agenda:',
    '  - Open dashboard, compare Week vs. Target',
    '  - Identify top 3 issues (CPA trending up? Response time slow?)',
    '  - Root cause analysis (creative fatigue? Audience saturation? QA issue?)',
    'Output: Action list for week (who fixes what by when)',
    'Document: Meeting notes in shared Google Doc'
])

add_heading('10.2 Phase 2 - Tuesday Channel Optimization (1 giờ, 2:00pm)', 2)
add_bullet_list([
    'Attendees: Digital Specialist, Content Creator',
    'Objective: Optimize campaigns based on Monday data',
    'Actions:',
    '  - Google Ads: Pause underperforming keywords, refresh bid strategy',
    '  - Meta: Swap creative, adjust audience, scale top performers',
    '  - Landing Page: A/B test form length, CTA copy, load time',
    '  - Funnel: Check conversion rates, identify bottlenecks',
    'Output: Optimization checklist signed off',
    'Tools: Birch, Google Ads, Meta Ads Manager, GHL'
])

add_heading('10.3 Phase 3 - Wednesday Pipeline Review (1.5 giờ, 10:00am)', 2)
add_bullet_list([
    'Attendees: Admissions Officers, Campus Leads, Marketing Manager',
    'Objective: Review every lead in pipeline, ensure no fall-through',
    'Process:',
    '  - Open GHL pipeline, sort by stage',
    '  - "New Lead" stage: Call all, qualify or close',
    '  - "Contacted" stage: Check SLA (should be 90% qualified by now)',
    '  - "Qualified" stage: Book tour or close',
    '  - "Booked" stage: Call morning-of for confirmation',
    '  - "Showed Tour" stage: Deliver offer same day',
    '  - "Offer" stage: Follow up every 2 days, create urgency',
    'Output: Updated GHL status for all leads, follow-up tasks assigned'
])

add_heading('10.4 Phase 4 - Thursday Content & Campaign Launch (2 giờ, 2:00pm)', 2)
add_bullet_list([
    'Attendees: Content Creator, Digital Specialist, Principal (optional)',
    'Objective: Launch new creative assets & campaigns for next week',
    'Tasks:',
    '  - Content: Review 3-5 new video/image assets (campus tour, testimonial, founder message)',
    '  - Ad Copy: Write & approve 5-10 new ad variations per platform',
    '  - Landing Page: Deploy new page (if any), test load & form',
    '  - Campaigns: Create & schedule new campaign sets (Google, Meta, TikTok)',
    '  - Email/SMS: Load next week nurture sequences',
    '  - Zalo OA: Queue rich message content',
    'Output: All new assets live & scheduled',
    'Quality check: Marketing Manager approves before go-live'
])

add_heading('10.5 Phase 5 - Friday Forecast & Planning (1 giờ, 3:00pm)', 2)
add_bullet_list([
    'Attendees: Marketing Manager, Campus Leads, Principal',
    'Objective: Forecast next week, anticipate challenges',
    'Items:',
    '  - Forecast: Predict leads, tours, enrollments for next week (based on trend)',
    '  - Events: Any Open Days, webinars, holiday impacts next week?',
    '  - Budget: Check YTD spend vs. budget, adjust allocations if needed',
    '  - Staffing: Are we staffed for forecasted volume? (Telesales, admissions)',
    '  - Challenges: What risks? (Competitor launch, platform outage, team absence)',
    '  - Next Week Priority: Top 3 focus areas',
    'Output: Weekly forecast memo, staffing plan, risk mitigation strategy'
])

add_page_break()

# Continue with remaining chapters (condensed for space)
add_heading('Chương 11: Campus Ownership & RACI', 1)
add_paragraph('Mỗi campus có Campus Lead chịu trách nhiệm KPI cuối cùng.')
add_heading('11.1 Campus Ownership Model', 2)
add_bullet_list([
    'Campus Lead (1 per campus): Chịu trách nhiệm # enrolled, tour show-up rate, offer conversion',
    'Campus Lead Reports To: Principal (local) + Marketing Manager (functional)',
    'KPI Accountability: # enrolled by end of Sept, CAC per student, parent satisfaction score',
    'Monthly Bonus: If campus hits 100% enrollment target'
])
add_heading('11.2 RACI Matrix', 2)
add_table([
    ['Task', 'Marketing', 'Admissions', 'Campus Lead', 'Principal'],
    ['Lead generation', 'R/A', 'S', 'S', 'S'],
    ['Campus tour quality', 'S', 'R/A', 'R/A', 'S'],
    ['Enrollment paperwork', 'S', 'R/A', 'S', 'S'],
    ['Budget decisions', 'R/A', 'S', 'S', 'A']
], [1.2, 1.2, 1.2, 1.2, 1.2])

add_page_break()

add_heading('Chương 12: Content Engine - Nội Dung Theo Funnel', 1)
add_paragraph('Content được tạo chủ yếu theo funnel stage, không làm content vì content.')
add_heading('12.1 Content Map By Funnel Stage', 2)
add_paragraph('TOFU (Awareness):')
add_bullet_list([
    'Video: "5 Lý Do Chọn Trường Việt Anh" (2-3 min)',
    'Blog: "Cách Chọn Trường Mầm Non Tốt" (1,500 words, SEO-optimized)',
    'Infographic: "Curriculum Comparison Chart" (visual)',
    'KPI: 500+ views/week per asset'
])
add_paragraph('MOFU (Consideration):')
add_bullet_list([
    'Testimonial Video: 3-4 parent/student videos (30-60 sec)',
    'School Tour Video: 5-7 min campus walkaround (4K)',
    'Comparison Guide: "VA vs. Other Schools" (PDF guide)',
    'FAQ Page: Q&A on curriculum, tuition, admission',
    'KPI: 1,000+ downloads/week, 25%+ email open rate'
])
add_paragraph('BOFU (Decision):')
add_bullet_list([
    'Offer Details: Pricing sheet, discount tiers, payment plans',
    'Enrollment Guide: Step-by-step "How to Enroll" checklist',
    'Parent Success Stories: 3 short case studies (500 words each)',
    'Deadline Reminder: Urgency-driven countdown email/SMS',
    'KPI: 60%+ offer → deposit conversion'
])

add_page_break()

# PHẦN D
add_heading('PHẦN D - CÔNG NGHỆ & TRACKING', 1)
add_page_break()

add_heading('Chương 13: Hệ Thống CRM Kép (GHL + Pancake)', 1)
add_heading('13.1 GoHighLevel (Pipeline Management)', 2)
add_bullet_list([
    'Role: Central pipeline hub, lead tracking, automation',
    'Features: Custom pipeline (8 stages), Contacts database, Workflows, Calendar, Reports',
    'Integration: Google Sheets (dashboard), Birch (ad platform)',
    'User Seats: Marketing Manager, Digital Specialist, Admissions Officers (5), Campus Leads (6)'
])
add_heading('13.2 Pancake CRM (Social Chat Management)', 2)
add_bullet_list([
    'Role: Facebook/Zalo message inbox consolidation',
    'Features: Chat inbox (unified), Lead distribution, Canned responses, Collaboration',
    'Integration: Webhooks to GHL (sync contact info, create tasks)',
    'User Seats: Admissions Officers (5), Support team (2)'
])
add_heading('13.3 Make.com Sync Workflow', 2)
add_bullet_list([
    'Sync Direction: Pancake → GHL (primary)',
    'Trigger: New chat message from Facebook/Zalo',
    'Action: Extract info, create/update GHL contact, create task, log interaction',
    'Reverse Sync: GHL stage change → Pancake note',
    'Error Handling: Slack notification if sync fails'
])

add_page_break()

add_heading('Chương 14: Funnel Chi Tiết & Landing Pages', 1)
add_paragraph('24 landing page funnels at forms.truongvietanh.com, each optimized per campaign type.')
add_heading('14.1 Funnel Structure Example (MN Program)', 2)
add_bullet_list([
    'Page 1 (Interest): Hero banner, 3 reasons, CTA "Learn More"',
    'Page 2 (Consideration): Curriculum, gallery, testimonials, CTA "Book Tour"',
    'Page 3 (Tour Booking): Form (name, email, phone, date), submit → GHL',
    'Page 4 (Confirmation): "Tour scheduled" message + calendar add',
    'Pixel Tracking: Page 1 → Interest, Page 2 → Consideration, Form → Lead'
])
add_heading('14.2 24 Funnels Breakdown', 2)
add_table([
    ['Funnel Type', 'Count', 'Examples'],
    ['Grade-based', '4', 'MN, TH, THCS, THPT'],
    ['Seasonal', '3', 'Spring, Summer, Fall intake'],
    ['Event-based', '4', 'Open Day, Webinar, Info Session'],
    ['Special Offer', '4', 'Early Bird, Sibling, Referral'],
    ['Retargeting', '5', 'Abandoned tour, Video watcher, Site visitor'],
    ['Campus-specific', '4', 'Gò Vấp, Bình Tân, Tây Ninh, An Giang'],
    ['TOTAL', '24', '']
], [1.5, 1.0, 3.5])

add_page_break()

add_heading('Chương 15: Analytics, GTM & Remarketing', 1)
add_heading('15.1 GTM Tag Setup (25 Tags)', 2)
add_bullet_list([
    'Google Analytics 4 (Page View, Lead)',
    'Google Ads Conversion (Lead form submit)',
    'Meta Pixel (Page View, Lead, Purchase)',
    'TikTok Pixel (Page View, Lead)',
    'Zalo Pixel (Page View, Lead)',
    'Facebook Lead Ads Pixel',
    'Hotjar (Session recording)',
    'CAPI setup (Server-side tracking)',
    'Plus 17 more tracking tags for comprehensive analytics'
])
add_heading('15.2 CAPI Setup (Server-Side Tracking)', 2)
add_bullet_list([
    'Purpose: Bypass ad blocker, improve match quality',
    'Events: ViewContent, Lead, Purchase, Schedule',
    'Implementation: Node.js webhook on backend',
    'Hashing: Phone, email, names hashed (SHA-256) before send'
])
add_heading('15.3 Remarketing Audiences (7 Active)', 2)
add_bullet_list([
    'Audience 1: Website Visitors (30 day window)',
    'Audience 2: Form Abandoners (14 day)',
    'Audience 3: Lead Database (180 day)',
    'Audience 4: Tour No-Shows (7 day - aggressive)',
    'Audience 5: Offer Viewers (10 day)',
    'Audience 6: Video Viewers (30 day)',
    'Exclusion: Enrolled students (remove immediately)'
])

add_page_break()

# PHẦN E
add_heading('PHẦN E - KIỂM SOÁT & CẢI TIẾN', 1)
add_page_break()

add_heading('Chương 16: QA & Mystery Call Program', 1)
add_heading('16.1 Mystery Call Program', 2)
add_bullet_list([
    'Frequency: 1 mystery call per campus per month (6 calls/month total)',
    'Evaluator: External QA vendor (không phải nội bộ)',
    'Scenarios: Incoming call, booking request, objection handling, followup call',
    'Recording: All calls recorded, shared with campus lead'
])
add_heading('16.2 Call Scoring Rubric (100 points)', 2)
add_table([
    ['Criteria', 'Points', 'Ideal Response'],
    ['Answer time < 5 sec', '10', 'Pick up immediately'],
    ['Warm greeting', '10', 'Use name, school name, cheerful tone'],
    ['Needs assessment Qs', '15', 'Ask 3+ open-end questions'],
    ['Product knowledge', '15', 'Articulate curriculum, facilities, pricing'],
    ['Objection handling', '15', 'Address concern, reframe, social proof'],
    ['Booking specificity', '15', 'Offer multiple date/time options'],
    ['Follow-up setup', '10', 'Confirm email, send calendar invite'],
    ['Professionalism', '5', 'No background noise, clear diction']
], [1.2, 0.8, 2.0])

add_page_break()

add_heading('Chương 17: Lost Reasons Tracking & Win-Back', 1)
add_heading('17.1 Lost Reason Categories', 2)
add_bullet_list([
    'Giá cao (Price objection)',
    'Xa (Location/distance)',
    'Chọn trường khác (Chose competitor)',
    'Chưa sẵn sàng (Decision not ready)',
    'Không liên lạc được (No response)',
    'Khác (Other)'
])
add_heading('17.2 Monthly Lost Analysis', 2)
add_paragraph('Every end of month:')
add_bullet_list([
    'Pull "Lost" leads by reason from GHL',
    'Identify top 3 reasons (usually 60-70% of losses)',
    'Create action plan per reason',
    'Execute improvements for next month'
])
add_heading('17.3 Win-Back Campaigns', 2)
add_bullet_list([
    'Target: Recoverable leads (lost < 30 days, showed tour)',
    'Campaign: 3-email sequence + phone call over 2 weeks',
    'Success Rate Target: 10-15% win-back rate',
    'Budget: Included in committed budget (telesales time)'
])

add_page_break()

add_heading('Chương 18: Offer & Deadline Management', 1)
add_heading('18.1 Early Bird Pricing Tiers', 2)
add_table([
    ['Tier', 'Dates', 'Discount', 'Tuition Reduction'],
    ['Tier 1 (Early Bird)', '26/03 - 30/04', '15%', 'Apply to full year'],
    ['Tier 2 (Priority)', '01/05 - 31/05', '10%', 'Apply to full year'],
    ['Tier 3 (Standard)', '01/06 - 31/08', '5%', 'Apply to half year'],
    ['Tier 4 (Last Minute)', '01/09 - 30/09', '0%', 'No discount'],
    ['Scholarship', 'All year', 'Varies', 'Merit/need-based']
], [1.5, 1.5, 1.5, 1.5])
add_heading('18.2 Urgency Creation Tools', 2)
add_bullet_list([
    'Limited Spots: "Only 8 spots left for Grade 1"',
    'Deadline Countdown: "Early bird pricing expires in 10 days"',
    'Exclusivity: "This offer only for tour participants"',
    'Scarcity: "Spot guarantees only with deposit by [DATE]"',
    'Social Proof: "15 families enrolled this week, interest is strong"'
])

add_page_break()

add_heading('Chương 19: Nhân Sự & Đào Tạo', 1)
add_heading('19.1 Team Structure (Tổng 16+ FTE)', 2)
add_table([
    ['Role', 'Count', 'Reports To', 'Key Responsibilities'],
    ['Marketing Manager', '1', 'Principal', 'Strategy, budget, KPI ownership'],
    ['Digital Specialist', '1', 'Marketing Mgr', 'Google Ads, Meta, analytics'],
    ['Content Creator', '1', 'Marketing Mgr', 'Video, blog, social media'],
    ['Admissions Officer (HQ)', '2', 'Marketing Mgr', 'Inbound leads, qualification'],
    ['Campus Lead', '6', 'Principal + Admissions Mgr', 'Tour coordination, campus KPI'],
    ['Admissions Officer (Campus)', '4-5', 'Campus Lead', 'Tour execution, offers'],
    ['Telesales', '2', 'Admissions Mgr', 'Outbound calling, closing'],
    ['TOTAL', '16-18', '', '']
], [1.5, 0.8, 1.5, 2.2])
add_heading('19.2 Training Plan (Pre-Launch & Ongoing)', 2)
add_paragraph('Pre-Launch Training (Weeks 1-2):')
add_bullet_list([
    'Product knowledge (curriculum, facilities, alumni outcomes)',
    'Sales skills (discovery, needs analysis, objection handling)',
    'CRM training (GHL navigation, lead management)',
    'Roleplay: Practice tours, handle objections'
])
add_paragraph('Ongoing Training (Monthly):')
add_bullet_list([
    'Sales coaching (review call recordings, 1-1 feedback)',
    'Product updates (new programs, pricing changes)',
    'Competitive analysis (how to compare with other schools)',
    'New campaign tactics (updated marketing approach)'
])
add_heading('19.3 Onboarding Checklist (For New Hires)', 2)
add_paragraph('Week 1:')
add_bullet_list([
    '[ ] Office setup (desk, computer, phone)',
    '[ ] System access (GHL, Pancake, Google Workspace)',
    '[ ] Welcome meeting with supervisor',
    '[ ] Product orientation (curriculum, campuses)',
    '[ ] Shadow existing team member (2 calls/tours)'
])
add_paragraph('Week 2-4:')
add_bullet_list([
    '[ ] CRM training (hands-on GHL)',
    '[ ] Sales training (objection handling workshop)',
    '[ ] Role-play evaluation (script practice)',
    '[ ] First independent call (with supervisor listening)',
    '[ ] Full responsibility with check-ins'
])

add_page_break()

# Conclusion
add_heading('KẾT LUẬN', 1)
add_paragraph('Kế hoạch tuyển sinh 2026 phiên bản 11 này biến đổi từ "marketing plan" sang "enrollment operating system" với 19 chapters chi tiết, từ nền tảng hệ thống (4 chapters) đến cấu trúc kênh & ngân sách (3 chapters), vận hành hàng ngày (4 chapters), công nghệ & tracking (3 chapters), và kiểm soát & cải tiến (4 chapters).', space_before=10, space_after=10)

add_paragraph('Mục tiêu rõ ràng: 1,300 học sinh, 13 tỷ VND, 28 tuần. Phương pháp: 8-tầng funnel, 4-trục kênh, SLA chặt chẽ, QA toàn diện. Khác biệt: Không chỉ "chạy quảng cáo" mà "vận hành hệ thống"—định nghĩa rõ quy trình, SOP, KPI, CRM, analytics. Thành công phụ thuộc vào execution, team alignment, và data-driven optimization tuần sau tuần.', space_before=10, space_after=10)

add_paragraph('Ngày bắt đầu: 26/03/2026. Deadline enrollment: 30/09/2026. Bắt đầu từ 72 giờ đầu tiên: Lock 20 quyết định chiến lược, xác nhận team, setup công nghệ. Sau đó, chạy hàng tuần theo 5-phase battleplan, track 12 metrics hàng ngày, cải tiến liên tục.', space_before=10, space_after=10)

# Save document
output_path = '/sessions/magical-youthful-faraday/mnt/truongvietanh.com/KH_Tuyen_Sinh_2026_V11_ChiTiet.docx'
doc.save(output_path)
print(f"DOCX generated successfully: {output_path}")

# Get file size
import os
file_size_mb = os.path.getsize(output_path) / (1024 * 1024)
print(f"File size: {file_size_mb:.2f} MB")
